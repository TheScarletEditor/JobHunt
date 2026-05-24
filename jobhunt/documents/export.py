from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches, RGBColor

from reportlab.lib.colors import HexColor
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate

from .model import (
    ResumeContent,
    SECTION_KIND_DETAIL, SECTION_KIND_LIST, SECTION_KIND_LINE,
    detect_section_kind,
)


def _escape(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def export_resume_docx(content: ResumeContent, path: str | Path) -> Path:
    out = Path(path)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    if content.name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(content.name)
        run.bold = True
        run.font.size = Pt(20)

    if content.contact:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("  ·  ".join(content.contact))
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if content.summary:
        doc.add_paragraph()
        p = doc.add_paragraph(content.summary)
        for run in p.runs:
            run.font.size = Pt(11)

    for s in content.sections:
        doc.add_paragraph()
        if s.title:
            p = doc.add_paragraph()
            run = p.add_run(s.title.upper())
            run.bold = True
            run.font.size = Pt(12)

        kind = detect_section_kind(s.title)

        if kind == SECTION_KIND_LIST:
            for item in s.items:
                skills = ", ".join(b for b in item.bullets if b)
                p = doc.add_paragraph()
                if item.header:
                    run = p.add_run(f"{item.header}: ")
                    run.bold = True
                    run.font.size = Pt(10)
                if skills:
                    run = p.add_run(skills)
                    run.font.size = Pt(10)
        elif kind == SECTION_KIND_LINE:
            for item in s.items:
                if not (item.header or item.subheader):
                    continue
                line = item.header or ""
                if item.subheader:
                    line = f"{line} · {item.subheader}" if line else item.subheader
                p = doc.add_paragraph(line, style="List Bullet")
                for run in p.runs:
                    run.font.size = Pt(10)
        else:  # SECTION_KIND_DETAIL
            for item in s.items:
                if item.header:
                    p = doc.add_paragraph()
                    run = p.add_run(item.header)
                    run.bold = True
                    run.font.size = Pt(11)
                if item.subheader:
                    p = doc.add_paragraph()
                    run = p.add_run(item.subheader)
                    run.italic = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                for bullet in item.bullets:
                    p = doc.add_paragraph(bullet, style="List Bullet")
                    for run in p.runs:
                        run.font.size = Pt(10)

    doc.save(str(out))
    return out


def export_cover_letter_docx(content: str, path: str | Path, profile: dict | None = None) -> Path:
    out = Path(path)
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    if profile:
        name = profile.get("legal_name") or profile.get("preferred_name") or ""
        if name:
            p = doc.add_paragraph()
            run = p.add_run(name)
            run.bold = True
            run.font.size = Pt(14)
        bits: list[str] = []
        for k in ("email", "phone", "linkedin_url"):
            if profile.get(k):
                bits.append(str(profile[k]))
        if bits:
            p = doc.add_paragraph()
            run = p.add_run(" · ".join(bits))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        doc.add_paragraph()

    for paragraph in (content or "").split("\n\n"):
        text = paragraph.strip()
        if text:
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(11)

    doc.save(str(out))
    return out


def export_resume_pdf(content: ResumeContent, path: str | Path) -> Path:
    out = Path(path)
    doc = SimpleDocTemplate(
        str(out), pagesize=LETTER,
        leftMargin=0.8 * inch, rightMargin=0.8 * inch,
        topMargin=0.7 * inch, bottomMargin=0.7 * inch,
    )

    name_style = ParagraphStyle(
        "Name", fontSize=20, fontName="Helvetica-Bold",
        alignment=TA_CENTER, spaceAfter=4,
    )
    contact_style = ParagraphStyle(
        "Contact", fontSize=10, fontName="Helvetica",
        alignment=TA_CENTER, spaceAfter=14, textColor=HexColor("#555555"),
    )
    summary_style = ParagraphStyle(
        "Summary", fontSize=11, fontName="Helvetica",
        alignment=TA_LEFT, spaceAfter=12, leading=14,
    )
    section_style = ParagraphStyle(
        "Section", fontSize=12, fontName="Helvetica-Bold",
        alignment=TA_LEFT, spaceBefore=10, spaceAfter=4,
        textColor=HexColor("#222222"),
    )
    item_header_style = ParagraphStyle(
        "ItemHeader", fontSize=11, fontName="Helvetica-Bold",
        alignment=TA_LEFT, spaceBefore=4, spaceAfter=0,
    )
    item_subheader_style = ParagraphStyle(
        "ItemSubheader", fontSize=10, fontName="Helvetica-Oblique",
        alignment=TA_LEFT, spaceAfter=4, textColor=HexColor("#666666"),
    )
    bullet_style = ParagraphStyle(
        "Bullet", fontSize=10, fontName="Helvetica",
        alignment=TA_LEFT, leading=13, leftIndent=14, spaceAfter=0,
    )

    story: list = []
    if content.name:
        story.append(Paragraph(_escape(content.name), name_style))
    if content.contact:
        story.append(Paragraph(_escape("  ·  ".join(content.contact)), contact_style))
    if content.summary:
        story.append(Paragraph(_escape(content.summary), summary_style))

    list_line_style = ParagraphStyle(
        "ListLine", fontSize=10, fontName="Helvetica",
        alignment=TA_LEFT, leading=13, spaceAfter=2,
    )

    for s in content.sections:
        title = (s.title or "").upper()
        if title:
            story.append(Paragraph(_escape(title), section_style))
        kind = detect_section_kind(s.title)

        if kind == SECTION_KIND_LIST:
            for item in s.items:
                skills = ", ".join(b for b in item.bullets if b)
                if item.header and skills:
                    body = f"<b>{_escape(item.header)}:</b> {_escape(skills)}"
                elif item.header:
                    body = f"<b>{_escape(item.header)}</b>"
                elif skills:
                    body = _escape(skills)
                else:
                    continue
                story.append(Paragraph(body, list_line_style))
        elif kind == SECTION_KIND_LINE:
            for item in s.items:
                if not (item.header or item.subheader):
                    continue
                line = _escape(item.header or "")
                if item.subheader:
                    sub = _escape(item.subheader)
                    line = f"{line} · <i>{sub}</i>" if line else f"<i>{sub}</i>"
                story.append(Paragraph(f"• {line}", bullet_style))
        else:  # SECTION_KIND_DETAIL
            for item in s.items:
                if item.header:
                    story.append(Paragraph(_escape(item.header), item_header_style))
                if item.subheader:
                    story.append(Paragraph(_escape(item.subheader), item_subheader_style))
                for bullet in item.bullets:
                    story.append(Paragraph(f"• {_escape(bullet)}", bullet_style))

    doc.build(story)
    return out


def export_cover_letter_pdf(content: str, path: str | Path, profile: dict | None = None) -> Path:
    out = Path(path)
    doc = SimpleDocTemplate(
        str(out), pagesize=LETTER,
        leftMargin=1 * inch, rightMargin=1 * inch,
        topMargin=1 * inch, bottomMargin=1 * inch,
    )

    name_style = ParagraphStyle(
        "Name", fontSize=14, fontName="Helvetica-Bold", spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", fontSize=10, fontName="Helvetica",
        textColor=HexColor("#666666"), spaceAfter=18,
    )
    body_style = ParagraphStyle(
        "Body", fontSize=11, fontName="Helvetica", leading=15, spaceAfter=11,
    )

    story: list = []
    if profile:
        name = profile.get("legal_name") or profile.get("preferred_name") or ""
        if name:
            story.append(Paragraph(_escape(name), name_style))
        bits = []
        for k in ("email", "phone", "linkedin_url"):
            if profile.get(k):
                bits.append(str(profile[k]))
        if bits:
            story.append(Paragraph(_escape(" · ".join(bits)), contact_style))

    for para in (content or "").split("\n\n"):
        text = para.strip()
        if text:
            story.append(
                Paragraph(_escape(text).replace("\n", "<br/>"), body_style)
            )

    doc.build(story)
    return out
