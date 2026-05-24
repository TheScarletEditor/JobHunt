from __future__ import annotations

import re
from pathlib import Path

from .model import ResumeContent, ResumeSection, ResumeItem


KNOWN_SECTIONS = {
    "summary", "objective", "profile", "about",
    "experience", "work experience", "professional experience", "employment",
    "employment history", "work history",
    "education", "academic background",
    "skills", "technical skills", "core skills", "competencies",
    "projects", "selected projects",
    "certifications", "licenses", "certifications & licenses",
    "awards", "honors", "honors & awards",
    "publications", "talks",
    "volunteer", "volunteer experience",
    "languages",
    "interests", "hobbies",
    "references",
}

BULLET_PREFIXES = ("•", "●", "▪", "■", "○", "·", "-", "*", "—", "–")
EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
PHONE_RE = re.compile(r"(?:\+?\d[\d\-\s().]{7,}\d)")
URL_RE = re.compile(r"https?://\S+|(?:linkedin|github)\.com/\S+")


def extract_raw_text(path: str | Path) -> str:
    """Return the raw text content of a resume file with no structuring."""
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(str(path))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            pass
    if ext == ".pdf":
        try:
            import pdfplumber
            with pdfplumber.open(str(path)) as pdf:
                texts = []
                for page in pdf.pages:
                    txt = page.extract_text() or ""
                    if txt.strip():
                        texts.append(txt)
                if texts:
                    return "\n\n".join(texts)
        except Exception:
            pass
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            return "\n\n".join((p.extract_text() or "") for p in reader.pages)
        except Exception:
            pass
    try:
        return path.read_text(encoding="utf-8", errors="ignore")
    except Exception:
        return ""


def parse_file(path: str | Path) -> ResumeContent:
    path = Path(path)
    ext = path.suffix.lower()
    if ext == ".docx":
        try:
            return _parse_docx(path)
        except Exception:
            pass
    if ext == ".pdf":
        try:
            return _parse_pdf(path)
        except Exception:
            pass
    return parse_text(path.read_text(encoding="utf-8", errors="ignore"))


def _parse_docx(path: Path) -> ResumeContent:
    import docx
    doc = docx.Document(str(path))
    lines: list[tuple[str, str]] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append(("blank", ""))
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if "heading" in style or "title" in style:
            kind = "heading"
        elif "list" in style or text[0] in BULLET_PREFIXES:
            kind = "bullet"
        else:
            kind = "text"
        lines.append((kind, _strip_bullet(text)))
    return _assemble(lines)


def _parse_pdf(path: Path) -> ResumeContent:
    text = ""
    try:
        import pdfplumber
        with pdfplumber.open(str(path)) as pdf:
            text = "\n".join((page.extract_text() or "") for page in pdf.pages)
    except Exception:
        text = ""
    if not text.strip():
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            text = "\n".join((page.extract_text() or "") for page in reader.pages)
        except Exception:
            text = ""
    return parse_text(text)


def parse_text(text: str) -> ResumeContent:
    raw_lines = text.replace("\r", "").split("\n")
    lines: list[tuple[str, str]] = []
    for raw in raw_lines:
        s = raw.strip()
        if not s:
            lines.append(("blank", ""))
            continue
        if s[0] in BULLET_PREFIXES:
            lines.append(("bullet", _strip_bullet(s)))
        elif _looks_like_section(s):
            lines.append(("heading", s))
        else:
            lines.append(("text", s))
    return _assemble(lines)


def _strip_bullet(s: str) -> str:
    while s and s[0] in BULLET_PREFIXES:
        s = s[1:].lstrip()
    return s


def _looks_like_section(s: str) -> bool:
    if len(s) > 60 or len(s.split()) > 5:
        return False
    lowered = s.strip(" :").lower()
    if lowered in KNOWN_SECTIONS:
        return True
    if s.isupper() and len(s) >= 3:
        return True
    return False


def _assemble(lines: list[tuple[str, str]]) -> ResumeContent:
    content = ResumeContent()

    i = 0
    n = len(lines)

    while i < n and lines[i][0] == "blank":
        i += 1

    if i < n and lines[i][0] in ("heading", "text"):
        content.name = lines[i][1]
        i += 1

    contact: list[str] = []
    while i < n and lines[i][0] == "text":
        line = lines[i][1]
        if EMAIL_RE.search(line) or PHONE_RE.search(line) or URL_RE.search(line) or "|" in line or "•" in line:
            contact.append(line)
            i += 1
        else:
            break
    content.contact = contact

    current_section: ResumeSection | None = None
    current_item: ResumeItem | None = None
    pending_summary: list[str] = []
    saw_first_heading = False

    while i < n:
        kind, text = lines[i]
        if kind == "blank":
            i += 1
            continue

        if kind == "heading":
            if not saw_first_heading and pending_summary:
                content.summary = " ".join(pending_summary).strip()
                pending_summary = []
            saw_first_heading = True
            current_section = ResumeSection(title=text.strip(" :"))
            content.sections.append(current_section)
            current_item = None
        elif kind == "bullet":
            if current_section is None:
                current_section = ResumeSection(title="")
                content.sections.append(current_section)
            if current_item is None:
                current_item = ResumeItem()
                current_section.items.append(current_item)
            current_item.bullets.append(text)
        else:
            if not saw_first_heading:
                pending_summary.append(text)
            else:
                if current_section is None:
                    current_section = ResumeSection(title="")
                    content.sections.append(current_section)
                if current_item is None or current_item.bullets:
                    current_item = ResumeItem(header=text)
                    current_section.items.append(current_item)
                elif not current_item.header:
                    current_item.header = text
                elif not current_item.subheader:
                    current_item.subheader = text
                else:
                    current_item = ResumeItem(header=text)
                    current_section.items.append(current_item)
        i += 1

    if pending_summary and not content.summary:
        content.summary = " ".join(pending_summary).strip()

    content.sections = [s for s in content.sections if s.title or s.items]
    return content
